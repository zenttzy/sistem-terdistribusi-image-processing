import csv
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parent.parent
OUTPUT = Path("/home/ubuntu/Laporan-UAS-Sistem-Terdistribusi-Hadi-Sanjaya-UNSRI.docx")
CSV_PATH = ROOT / "reports/aggregate/comparison.csv"
CHARTS = ROOT / "reports/aggregate"


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, color=None, size=8)
        set_cell_shading(table.rows[0].cells[index], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value, size=8)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    document.add_paragraph()
    return table


def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(5)
    return paragraph


def add_caption(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def add_bullet(document, text):
    document.add_paragraph(text, style="List Bullet")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def add_toc(document):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Klik kanan lalu pilih Update Field untuk menampilkan daftar isi."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def read_rows():
    with CSV_PATH.open(newline="", encoding="utf-8") as file:
        return list(csv.DictReader(file))


def main():
    rows = read_rows()
    rows_300 = [row for row in rows if row["total_images"] == "300"]
    sequential_300 = next(row for row in rows_300 if row["mode"] == "sequential")
    distributed_300 = [row for row in rows_300 if row["mode"] == "distributed"]
    best = min(distributed_300, key=lambda row: float(row["wall_time_seconds"]))

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(12)
    styles["Normal"].paragraph_format.line_spacing = 1.5
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[style_name].font.name = "Times New Roman"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_page_number(section.footer.paragraphs[0])

    # Sampul mengikuti struktur umum laporan tugas Unsri.
    for _ in range(2):
        document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LAPORAN TUGAS AKHIR SEMESTER\nMATA KULIAH SISTEM TERDISTRIBUSI")
    r.bold = True
    r.font.size = Pt(16)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("IMPLEMENTASI DAN ANALISIS KINERJA PEMROSESAN CITRA DIGITAL\nPARALEL DAN TERDISTRIBUSI MENGGUNAKAN ARSITEKTUR MASTER-WORKER")
    r.bold = True
    r.font.size = Pt(14)
    document.add_paragraph()
    logo = document.add_table(rows=1, cols=1)
    logo.alignment = WD_TABLE_ALIGNMENT.CENTER
    logo_cell = logo.cell(0, 0)
    set_cell_text(logo_cell, "TEMPATKAN LOGO RESMI\nUNIVERSITAS SRIWIJAYA\nDI SINI", bold=True, size=10)
    logo_cell.width = Cm(5)
    logo_cell.height = Cm(5)
    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Disusun oleh:\n\n").font.size = Pt(12)
    identity = p.add_run("HADI SANJAYA\n09011282328112")
    identity.bold = True
    identity.font.size = Pt(12)
    p.add_run("\n\nDosen Pengampu: ................................................").font.size = Pt(11)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    institution = p.add_run("PROGRAM STUDI SISTEM KOMPUTER\nFAKULTAS ILMU KOMPUTER\nUNIVERSITAS SRIWIJAYA\n2026")
    institution.bold = True
    institution.font.size = Pt(12)
    document.add_page_break()

    add_heading(document, "LEMBAR PENGESAHAN", 1)
    document.add_paragraph("Laporan UAS berjudul “Implementasi dan Analisis Kinerja Pemrosesan Citra Digital Paralel dan Terdistribusi Menggunakan Arsitektur Master-Worker” disusun oleh:")
    add_table(document, ["Identitas", "Keterangan"], [["Nama", "Hadi Sanjaya"], ["NIM", "09011282328112"], ["Program Studi", "Sistem Komputer"], ["Fakultas", "Ilmu Komputer"], ["Universitas", "Universitas Sriwijaya"]], [5, 10])
    document.add_paragraph("Telah diperiksa dan disetujui sebagai laporan tugas UAS Mata Kuliah Sistem Terdistribusi.")
    document.add_paragraph("\nIndralaya, ................................ 2026\n\n\n\nDosen Pengampu,\n\n\n\n(................................................)\nNIP. ...........................................")
    document.add_page_break()

    add_heading(document, "KATA PENGANTAR", 1)
    document.add_paragraph("Puji syukur penulis panjatkan kepada Tuhan Yang Maha Esa karena laporan UAS Sistem Terdistribusi ini dapat diselesaikan. Laporan ini membahas perancangan, implementasi, pengujian, dan analisis kinerja sistem pemrosesan citra digital paralel dan terdistribusi menggunakan Python, OpenCV, Celery, dan Redis pada Ubuntu Server.")
    document.add_paragraph("Penulis menyampaikan terima kasih kepada dosen pengampu Mata Kuliah Sistem Terdistribusi, Fakultas Ilmu Komputer Universitas Sriwijaya, atas arahan dan materi yang diberikan. Penulis menyadari laporan ini masih dapat dikembangkan, khususnya melalui pengujian pada beberapa node fisik dan dataset citra nyata yang lebih besar.")
    document.add_paragraph("Indralaya, Juli 2026\n\nPenulis,\n\n\nHadi Sanjaya")
    document.add_page_break()

    add_heading(document, "Abstrak", 1)
    document.add_paragraph(
        "Pemrosesan citra dalam jumlah besar membutuhkan waktu komputasi yang cukup tinggi apabila dilakukan secara sequential. "
        "Project ini mengimplementasikan sistem pemrosesan citra digital paralel dan terdistribusi menggunakan arsitektur master-worker. "
        "Master membagi setiap gambar menjadi task, Redis digunakan sebagai message broker, dan Celery worker menjalankan pipeline OpenCV "
        "berupa resize, grayscale, Gaussian blur, dan Canny edge detection. Pengujian dilakukan pada Ubuntu Server dengan dua CPU core "
        "menggunakan dataset 300 gambar berukuran 1920x1080. Hasil terbaik diperoleh dengan dua worker, yaitu waktu 10,694 detik, "
        "throughput 28,05 gambar per detik, speedup 1,21, dan efficiency 60,74%."
    )
    document.add_paragraph("Kata kunci: sistem terdistribusi, pemrosesan citra, OpenCV, Celery, Redis, master-worker.")

    add_heading(document, "ABSTRACT", 1)
    document.add_paragraph("Digital image processing in large batches requires significant computation time when executed sequentially. This project implements a parallel and distributed image-processing system using a master-worker architecture. The master submits one task for each image, Redis acts as the message broker and result backend, and Celery workers execute an OpenCV pipeline consisting of resize, grayscale conversion, Gaussian blur, and Canny edge detection. The experiment used 300 Full HD images on Ubuntu Server with two allocated CPU cores. The best two-worker execution completed in 10.694 seconds with a throughput of 28.05 images per second, a speedup of 1.21, and an efficiency of 60.74 percent. The results demonstrate measurable acceleration while also showing the effect of messaging, scheduling, I/O, and resource-contention overhead.")
    document.add_paragraph("Keywords: distributed system, image processing, OpenCV, Celery, Redis, master-worker.")

    add_heading(document, "Daftar Isi", 1)
    document.add_paragraph("Pada Microsoft Word, klik kanan bagian ini lalu pilih Update Field untuk memperbarui daftar isi otomatis.")
    add_toc(document)
    add_heading(document, "Daftar Tabel", 1)
    document.add_paragraph("Tabel 1. Spesifikasi lingkungan pengujian\nTabel 2. Hasil benchmark dataset 300 gambar")
    add_heading(document, "Daftar Gambar", 1)
    document.add_paragraph("Gambar 1. Perbandingan waktu eksekusi\nGambar 2. Perbandingan throughput")
    document.add_page_break()

    add_heading(document, "BAB I PENDAHULUAN", 1)
    add_heading(document, "1.1 Latar Belakang", 2)
    document.add_paragraph(
        "Data citra digital terus bertambah dan sering membutuhkan operasi pemrosesan yang sama pada banyak file. "
        "Pemrosesan sequential membuat seluruh pekerjaan berjalan pada satu alur sehingga waktu penyelesaian meningkat. "
        "Sistem terdistribusi dapat membagi pekerjaan ke beberapa worker melalui antrean task sehingga pekerjaan independen dapat dikerjakan secara konkuren."
    )
    document.add_paragraph("Sebagai analogi terhadap distributed log analyzer yang membagi banyak file log untuk diproses worker, project ini membagi sekumpulan gambar menjadi task independen. Pendekatan per gambar dipilih karena mudah didistribusikan, memiliki isolasi kegagalan yang baik, dan tidak membutuhkan sinkronisasi piksel antarpotongan seperti pada pembagian satu citra menjadi beberapa tile.")
    document.add_paragraph("Ubuntu Server digunakan sebagai lingkungan implementasi karena aplikasi bersifat lintas distribusi Linux. Perbedaan dengan AlmaLinux terutama berada pada package manager dan pengelolaan service, sedangkan kode Python, OpenCV, Celery, dan Redis tetap sama.")
    add_heading(document, "1.2 Rumusan Masalah", 2)
    for text in [
        "Bagaimana merancang sistem pemrosesan citra dengan arsitektur master-worker?",
        "Bagaimana membagi task gambar menggunakan Redis dan Celery?",
        "Bagaimana pengaruh dua worker terhadap waktu, throughput, speedup, dan efficiency?",
        "Bagaimana sistem menangani task gambar yang gagal atau worker yang terhenti?",
    ]:
        add_bullet(document, text)
    add_heading(document, "1.3 Tujuan", 2)
    for text in [
        "Mengimplementasikan pemrosesan citra digital secara sequential dan distributed.",
        "Menerapkan komunikasi task menggunakan Redis dan Celery.",
        "Mengukur kinerja berdasarkan waktu eksekusi, throughput, speedup, dan efficiency.",
        "Menyediakan laporan hasil eksperimen yang dapat direproduksi.",
    ]:
        add_bullet(document, text)
    add_heading(document, "1.4 Batasan Masalah", 2)
    for text in [
        "Pengujian dilakukan pada satu Ubuntu Server dengan beberapa proses worker.",
        "Pipeline citra terdiri dari resize, grayscale, Gaussian blur, dan Canny edge detection.",
        "Redis berfungsi sebagai broker dan result backend Celery.",
        "GPU, deep learning, dan pembagian tile dalam satu gambar tidak dibahas.",
    ]:
        add_bullet(document, text)
    add_heading(document, "1.5 Manfaat", 2)
    document.add_paragraph("Project memberikan contoh penerapan konsep message broker, distributed task queue, worker concurrency, fault tolerance, dan evaluasi performa. Hasil implementasi juga dapat menjadi fondasi untuk pengembangan layanan batch processing, klasifikasi citra, pengolahan dokumen hasil pemindaian, dan analisis citra ilmiah.")
    add_heading(document, "1.6 Sistematika Penulisan", 2)
    document.add_paragraph("BAB I membahas pendahuluan. BAB II menjelaskan teori pendukung. BAB III menguraikan analisis kebutuhan dan perancangan. BAB IV menjelaskan implementasi, pengujian, serta analisis hasil. BAB V menyajikan kesimpulan dan saran.")

    add_heading(document, "BAB II LANDASAN TEORI", 1)
    for title, body in [
        ("2.1 Sistem Terdistribusi", "Sistem terdistribusi adalah kumpulan komputer atau proses yang bekerja sama melalui jaringan untuk menyelesaikan pekerjaan. Pada project ini, konsep tersebut disimulasikan menggunakan master dan beberapa proses worker."),
        ("2.2 Pemrosesan Paralel", "Pemrosesan paralel menjalankan beberapa pekerjaan secara bersamaan. Karena setiap gambar dapat diproses secara independen, dataset cocok dibagi menjadi task per gambar."),
        ("2.3 Arsitektur Master-Worker", "Master bertugas membuat dan mengirim pekerjaan, sedangkan worker mengambil pekerjaan dan mengembalikan hasil. Model ini memudahkan pembagian beban dan penambahan worker."),
        ("2.4 Redis dan Celery", "Redis menyediakan antrean dan penyimpanan hasil sementara. Celery menyediakan abstraksi task queue, worker, retry, acknowledgement, dan monitoring."),
        ("2.5 OpenCV", "OpenCV digunakan untuk membaca, mengubah ukuran, mengonversi warna, mengurangi noise, dan mendeteksi tepi pada citra."),
        ("2.6 Metrik Kinerja", "Speedup = waktu sequential / waktu parallel. Efficiency = speedup / jumlah worker x 100%. Throughput = jumlah gambar berhasil / waktu total."),
    ]:
        add_heading(document, title, 2)
        document.add_paragraph(body)

    add_heading(document, "BAB III PERANCANGAN SISTEM", 1)
    add_heading(document, "3.1 Analisis Kebutuhan", 2)
    add_heading(document, "3.1.1 Kebutuhan Fungsional", 3)
    for text in ["Sistem membaca banyak gambar dari direktori dataset.", "Sistem menyediakan mode sequential dan distributed.", "Master membentuk satu task untuk setiap gambar.", "Worker menjalankan pipeline OpenCV dan menyimpan output.", "Sistem mencatat task berhasil, gagal, durasi, task ID, dan worker.", "Sistem menghasilkan JSON, CSV, dan grafik perbandingan."]:
        add_bullet(document, text)
    add_heading(document, "3.1.2 Kebutuhan Nonfungsional", 3)
    for text in ["Aplikasi berjalan pada Ubuntu Server tanpa GUI.", "Proses dapat direproduksi melalui command-line interface.", "Worker memiliki nama unik agar dapat dimonitor.", "Dataset dan output tidak disimpan dalam Git karena ukurannya besar.", "Sistem mempertahankan portabilitas ke AlmaLinux."]:
        add_bullet(document, text)
    add_heading(document, "3.2 Arsitektur", 2)
    document.add_paragraph("Arsitektur project menggunakan alur berikut:")
    architecture = [
        "Master CLI: membaca dataset, mengirim task, mengumpulkan hasil, dan membuat laporan.",
        "Redis: menyimpan antrean task dan hasil Celery.",
        "Worker 1 dan Worker 2: menjalankan pipeline OpenCV.",
        "Output: menyimpan gambar hasil edge detection.",
        "Reports: menyimpan summary JSON, CSV, dan grafik evaluasi.",
    ]
    for text in architecture:
        add_bullet(document, text)
    add_heading(document, "3.3 Alur Pemrosesan", 2)
    for index, text in enumerate(["Master menemukan file gambar.", "Setiap gambar diubah menjadi satu task.", "Redis memasukkan task ke antrean.", "Worker mengambil dan memproses task.", "Worker menyimpan hasil dan mengembalikan metadata.", "Master menulis laporan eksperimen."], 1):
        document.add_paragraph(f"{index}. {text}")
    add_heading(document, "3.4 Struktur Direktori", 2)
    document.add_paragraph("app/ berisi kode aplikasi; scripts/ berisi utilitas; dataset/ berisi input; output/ berisi hasil; reports/ berisi laporan; docs/ berisi dokumentasi.")
    add_heading(document, "3.5 Fault Tolerance", 2)
    document.add_paragraph("Task menggunakan late acknowledgement, prefetch satu task, dan retry otomatis untuk error I/O. Jika worker terhenti sebelum acknowledgement, task dapat diproses kembali oleh worker yang tersedia.")

    add_heading(document, "BAB IV IMPLEMENTASI DAN PENGUJIAN", 1)
    add_heading(document, "4.1 Lingkungan Pengujian", 2)
    add_table(document, ["Komponen", "Spesifikasi"], [["OS", "Ubuntu 24.04.4 LTS"], ["CPU", "AMD EPYC 7K62, 2 core dialokasikan"], ["RAM", "1,9 GiB"], ["Swap", "1,9 GiB"], ["Dataset", "300 gambar, 1920x1080"], ["Worker", "2 proses, concurrency 1"]], [5, 10])
    add_heading(document, "4.2 Perintah Pengujian", 2)
    document.add_paragraph("Instalasi dan eksekusi utama dilakukan dengan perintah berikut:")
    for command in ["./scripts/setup-ubuntu.sh", "redis-cli ping", "./scripts/start_worker.sh worker1 1", "./scripts/start_worker.sh worker2 1", "python -m app.cli --mode sequential --workers 1", "python -m app.cli --mode distributed --workers 2 --timeout 3600", "python scripts/aggregate_reports.py"]:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(command)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    add_heading(document, "4.3 Implementasi Pipeline", 2)
    document.add_paragraph("Pipeline dimulai dengan cv2.imread untuk membaca citra. Lebar gambar dibatasi maksimal 1920 piksel dengan rasio aspek dipertahankan. Citra kemudian dikonversi dari BGR ke grayscale untuk mengurangi dimensi kanal. Gaussian blur kernel 11x11 diterapkan untuk mengurangi noise sebelum Canny edge detection menggunakan threshold 50 dan 150. Output disimpan sebagai PNG agar hasil tepi tidak mengalami kompresi lossy tambahan.")
    add_heading(document, "4.4 Implementasi Distribusi Task", 2)
    document.add_paragraph("Celery dikonfigurasi menggunakan Redis sebagai broker dan result backend. task_acks_late memastikan acknowledgement diberikan setelah task selesai, worker_prefetch_multiplier bernilai satu membantu pemerataan task, dan result_expires membatasi masa penyimpanan hasil. Master menggunakan Celery group untuk mengirim seluruh task dan menunggu hasilnya.")
    add_heading(document, "4.5 Hasil Benchmark", 2)
    table_rows = []
    for row in rows_300:
        table_rows.append([row["mode"], row["workers"], row["wall_time_seconds"][:6], row["throughput_images_per_second"][:6], row["speedup"][:6], row["efficiency_percent"][:6] + "%"])
    add_table(document, ["Mode", "Worker", "Waktu (s)", "Throughput", "Speedup", "Efficiency"], table_rows, [3.5, 2, 3, 4, 3, 3])
    document.add_paragraph(f"Hasil terbaik adalah distributed dengan dua worker, waktu {best['wall_time_seconds']} detik, throughput {best['throughput_images_per_second']} gambar/detik, speedup {best['speedup']}, dan efficiency {best['efficiency_percent']}%.")
    add_heading(document, "4.6 Grafik Hasil", 2)
    for image, caption in [("execution-time.png", "Gambar 1. Perbandingan waktu eksekusi"), ("throughput.png", "Gambar 2. Perbandingan throughput")]:
        path = CHARTS / image
        if path.exists():
            document.add_picture(str(path), width=Inches(6.1))
            add_caption(document, caption)
    add_heading(document, "4.7 Analisis", 2)
    document.add_paragraph("Distributed dua worker memberikan waktu tercepat pada pengujian 300 gambar. Percepatan belum linear karena seluruh worker berjalan pada satu node dengan dua CPU core. Overhead Redis, serialisasi task, scheduling Celery, dan I/O gambar ikut memengaruhi hasil. Temuan ini menunjukkan pentingnya menguji ukuran dataset yang cukup besar dan membandingkan jumlah worker dengan resource CPU yang tersedia.")
    add_heading(document, "4.8 Pengujian Fault Tolerance", 2)
    document.add_paragraph("Saat salah satu koneksi Termius terputus, proses worker lama tetap dapat berjalan di server. Ketika worker baru dijalankan menggunakan nama yang sama, Celery mendeteksi DuplicateNodenameWarning. Pengujian ini menunjukkan pentingnya hostname worker yang unik. Setelah proses worker duplikat dihentikan dan worker1 serta worker2 dijalankan ulang, status kembali menunjukkan dua node online.")
    add_heading(document, "4.9 Dokumentasi Screenshot", 2)
    document.add_paragraph("Screenshot terminal dapat ditempatkan pada bagian ini: spesifikasi server, Redis PONG, dua worker online, log task worker, master CLI, monitoring top, summary.json, comparison.csv, serta gambar input dan output.")

    add_heading(document, "BAB V PENUTUP", 1)
    add_heading(document, "5.1 Kesimpulan", 2)
    document.add_paragraph("Project berhasil menerapkan pemrosesan citra digital paralel dan terdistribusi menggunakan arsitektur master-worker, Redis, Celery, dan OpenCV pada Ubuntu Server. Sebanyak 300 gambar berhasil diproses tanpa task gagal. Konfigurasi distributed dua worker memberikan hasil terbaik dengan speedup 1,21 dan efficiency 60,74% pada eksperimen yang tercatat.")
    add_heading(document, "5.2 Saran", 2)
    for text in ["Melakukan pengujian pada beberapa server fisik.", "Menggunakan shared storage atau object storage untuk multi-node.", "Menambahkan dashboard Flower untuk monitoring.", "Menggunakan dataset citra nyata yang lebih beragam.", "Menguji GPU atau algoritma pemrosesan yang lebih berat."]:
        add_bullet(document, text)

    add_heading(document, "DAFTAR PUSTAKA", 1)
    for text in ["OpenCV Documentation. Image Processing Module. https://docs.opencv.org/", "Celery Documentation. Distributed Task Queue. https://docs.celeryq.dev/", "Redis Documentation. In-memory Data Store. https://redis.io/docs/", "Universitas Sriwijaya. Pedoman Umum Penulisan Karya Tulis Ilmiah. https://unsri.ac.id/", "Repository project. https://github.com/zenttzy/sistem-terdistribusi-image-processing"]:
        document.add_paragraph(text)

    document.add_heading("LAMPIRAN A — Struktur Project", level=1)
    document.add_paragraph("Repository: https://github.com/zenttzy/sistem-terdistribusi-image-processing")
    document.add_paragraph("Dokumentasi teks: docs/hasil-pengujian.txt")
    document.add_paragraph("Kode utama: app/cli.py, app/tasks.py, app/image_processor.py, app/reporting.py")
    document.add_paragraph("Grafik: reports/aggregate/execution-time.png dan reports/aggregate/throughput.png")

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
