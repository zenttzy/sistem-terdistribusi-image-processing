import csv
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parent.parent
OUTPUT = Path("/home/ubuntu/Laporan-UAS-Distributed-Image-Processing.docx")
CSV_PATH = ROOT / "reports/aggregate/comparison.csv"
CHARTS = ROOT / "reports/aggregate"


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, color=None, size=8)
        set_cell_shading(table.rows[0].cells[index], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value, size=8)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    document.add_paragraph()
    return table


def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(5)
    return paragraph


def add_caption(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def add_bullet(document, text):
    document.add_paragraph(text, style="List Bullet")


def read_rows():
    with CSV_PATH.open(newline="", encoding="utf-8") as file:
        return list(csv.DictReader(file))


def main():
    rows = read_rows()
    rows_300 = [row for row in rows if row["total_images"] == "300"]
    sequential_300 = next(row for row in rows_300 if row["mode"] == "sequential")
    distributed_300 = [row for row in rows_300 if row["mode"] == "distributed"]
    best = min(distributed_300, key=lambda row: float(row["wall_time_seconds"]))

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(12)
    styles["Normal"].paragraph_format.line_spacing = 1.5
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[style_name].font.name = "Times New Roman"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # Cover page.
    for _ in range(4):
        document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LAPORAN UAS\nSISTEM TERDISTRIBUSI")
    r.bold = True
    r.font.size = Pt(16)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("IMPLEMENTASI PEMROSESAN CITRA DIGITAL\nPARALEL DAN TERDISTRIBUSI")
    r.bold = True
    r.font.size = Pt(14)
    for _ in range(5):
        document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Disusun untuk memenuhi tugas UAS\nMata Kuliah Sistem Terdistribusi\n\n").font.size = Pt(12)
    p.add_run("Ubuntu Server • Python • OpenCV • Celery • Redis").italic = True
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("2026").bold = True
    document.add_page_break()

    add_heading(document, "Abstrak", 1)
    document.add_paragraph(
        "Pemrosesan citra dalam jumlah besar membutuhkan waktu komputasi yang cukup tinggi apabila dilakukan secara sequential. "
        "Project ini mengimplementasikan sistem pemrosesan citra digital paralel dan terdistribusi menggunakan arsitektur master-worker. "
        "Master membagi setiap gambar menjadi task, Redis digunakan sebagai message broker, dan Celery worker menjalankan pipeline OpenCV "
        "berupa resize, grayscale, Gaussian blur, dan Canny edge detection. Pengujian dilakukan pada Ubuntu Server dengan dua CPU core "
        "menggunakan dataset 300 gambar berukuran 1920x1080. Hasil terbaik diperoleh dengan dua worker, yaitu waktu 10,694 detik, "
        "throughput 28,05 gambar per detik, speedup 1,21, dan efficiency 60,74%."
    )
    document.add_paragraph("Kata kunci: sistem terdistribusi, pemrosesan citra, OpenCV, Celery, Redis, master-worker.")

    add_heading(document, "Daftar Isi", 1)
    document.add_paragraph("Pada Microsoft Word, klik kanan bagian ini lalu pilih Update Field untuk memperbarui daftar isi otomatis.")
    toc = OxmlElement("w:sdt")
    tocPr = OxmlElement("w:sdtPr")
    docPartObj = OxmlElement("w:docPartObj")
    docPartGallery = OxmlElement("w:docPartGallery")
    docPartGallery.set(qn("w:val"), "Table of Contents")
    docPartObj.append(docPartGallery)
    tocPr.append(docPartObj)
    toc.append(tocPr)
    document._body._body.append(toc)
    document.add_page_break()

    add_heading(document, "BAB I PENDAHULUAN", 1)
    add_heading(document, "1.1 Latar Belakang", 2)
    document.add_paragraph(
        "Data citra digital terus bertambah dan sering membutuhkan operasi pemrosesan yang sama pada banyak file. "
        "Pemrosesan sequential membuat seluruh pekerjaan berjalan pada satu alur sehingga waktu penyelesaian meningkat. "
        "Sistem terdistribusi dapat membagi pekerjaan ke beberapa worker melalui antrean task sehingga pekerjaan independen dapat dikerjakan secara konkuren."
    )
    add_heading(document, "1.2 Rumusan Masalah", 2)
    for text in [
        "Bagaimana merancang sistem pemrosesan citra dengan arsitektur master-worker?",
        "Bagaimana membagi task gambar menggunakan Redis dan Celery?",
        "Bagaimana pengaruh dua worker terhadap waktu, throughput, speedup, dan efficiency?",
        "Bagaimana sistem menangani task gambar yang gagal atau worker yang terhenti?",
    ]:
        add_bullet(document, text)
    add_heading(document, "1.3 Tujuan", 2)
    for text in [
        "Mengimplementasikan pemrosesan citra digital secara sequential dan distributed.",
        "Menerapkan komunikasi task menggunakan Redis dan Celery.",
        "Mengukur kinerja berdasarkan waktu eksekusi, throughput, speedup, dan efficiency.",
        "Menyediakan laporan hasil eksperimen yang dapat direproduksi.",
    ]:
        add_bullet(document, text)
    add_heading(document, "1.4 Batasan Masalah", 2)
    for text in [
        "Pengujian dilakukan pada satu Ubuntu Server dengan beberapa proses worker.",
        "Pipeline citra terdiri dari resize, grayscale, Gaussian blur, dan Canny edge detection.",
        "Redis berfungsi sebagai broker dan result backend Celery.",
        "GPU, deep learning, dan pembagian tile dalam satu gambar tidak dibahas.",
    ]:
        add_bullet(document, text)

    add_heading(document, "BAB II LANDASAN TEORI", 1)
    for title, body in [
        ("2.1 Sistem Terdistribusi", "Sistem terdistribusi adalah kumpulan komputer atau proses yang bekerja sama melalui jaringan untuk menyelesaikan pekerjaan. Pada project ini, konsep tersebut disimulasikan menggunakan master dan beberapa proses worker."),
        ("2.2 Pemrosesan Paralel", "Pemrosesan paralel menjalankan beberapa pekerjaan secara bersamaan. Karena setiap gambar dapat diproses secara independen, dataset cocok dibagi menjadi task per gambar."),
        ("2.3 Arsitektur Master-Worker", "Master bertugas membuat dan mengirim pekerjaan, sedangkan worker mengambil pekerjaan dan mengembalikan hasil. Model ini memudahkan pembagian beban dan penambahan worker."),
        ("2.4 Redis dan Celery", "Redis menyediakan antrean dan penyimpanan hasil sementara. Celery menyediakan abstraksi task queue, worker, retry, acknowledgement, dan monitoring."),
        ("2.5 OpenCV", "OpenCV digunakan untuk membaca, mengubah ukuran, mengonversi warna, mengurangi noise, dan mendeteksi tepi pada citra."),
        ("2.6 Metrik Kinerja", "Speedup = waktu sequential / waktu parallel. Efficiency = speedup / jumlah worker x 100%. Throughput = jumlah gambar berhasil / waktu total."),
    ]:
        add_heading(document, title, 2)
        document.add_paragraph(body)

    add_heading(document, "BAB III PERANCANGAN SISTEM", 1)
    add_heading(document, "3.1 Arsitektur", 2)
    document.add_paragraph("Arsitektur project menggunakan alur berikut:")
    architecture = [
        "Master CLI: membaca dataset, mengirim task, mengumpulkan hasil, dan membuat laporan.",
        "Redis: menyimpan antrean task dan hasil Celery.",
        "Worker 1 dan Worker 2: menjalankan pipeline OpenCV.",
        "Output: menyimpan gambar hasil edge detection.",
        "Reports: menyimpan summary JSON, CSV, dan grafik evaluasi.",
    ]
    for text in architecture:
        add_bullet(document, text)
    add_heading(document, "3.2 Alur Pemrosesan", 2)
    for index, text in enumerate(["Master menemukan file gambar.", "Setiap gambar diubah menjadi satu task.", "Redis memasukkan task ke antrean.", "Worker mengambil dan memproses task.", "Worker menyimpan hasil dan mengembalikan metadata.", "Master menulis laporan eksperimen."], 1):
        document.add_paragraph(f"{index}. {text}")
    add_heading(document, "3.3 Struktur Direktori", 2)
    document.add_paragraph("app/ berisi kode aplikasi; scripts/ berisi utilitas; dataset/ berisi input; output/ berisi hasil; reports/ berisi laporan; docs/ berisi dokumentasi.")
    add_heading(document, "3.4 Fault Tolerance", 2)
    document.add_paragraph("Task menggunakan late acknowledgement, prefetch satu task, dan retry otomatis untuk error I/O. Jika worker terhenti sebelum acknowledgement, task dapat diproses kembali oleh worker yang tersedia.")

    add_heading(document, "BAB IV IMPLEMENTASI DAN PENGUJIAN", 1)
    add_heading(document, "4.1 Lingkungan Pengujian", 2)
    add_table(document, ["Komponen", "Spesifikasi"], [["OS", "Ubuntu 24.04.4 LTS"], ["CPU", "AMD EPYC 7K62, 2 core dialokasikan"], ["RAM", "1,9 GiB"], ["Swap", "1,9 GiB"], ["Dataset", "300 gambar, 1920x1080"], ["Worker", "2 proses, concurrency 1"]], [5, 10])
    add_heading(document, "4.2 Perintah Pengujian", 2)
    document.add_paragraph("Instalasi dan eksekusi utama dilakukan dengan perintah berikut:")
    for command in ["./scripts/setup-ubuntu.sh", "redis-cli ping", "./scripts/start_worker.sh worker1 1", "./scripts/start_worker.sh worker2 1", "python -m app.cli --mode sequential --workers 1", "python -m app.cli --mode distributed --workers 2 --timeout 3600", "python scripts/aggregate_reports.py"]:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(command)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    add_heading(document, "4.3 Hasil Benchmark", 2)
    table_rows = []
    for row in rows_300:
        table_rows.append([row["mode"], row["workers"], row["wall_time_seconds"][:6], row["throughput_images_per_second"][:6], row["speedup"][:6], row["efficiency_percent"][:6] + "%"])
    add_table(document, ["Mode", "Worker", "Waktu (s)", "Throughput", "Speedup", "Efficiency"], table_rows, [3.5, 2, 3, 4, 3, 3])
    document.add_paragraph(f"Hasil terbaik adalah distributed dengan dua worker, waktu {best['wall_time_seconds']} detik, throughput {best['throughput_images_per_second']} gambar/detik, speedup {best['speedup']}, dan efficiency {best['efficiency_percent']}%.")
    add_heading(document, "4.4 Grafik Hasil", 2)
    for image, caption in [("execution-time.png", "Gambar 1. Perbandingan waktu eksekusi"), ("throughput.png", "Gambar 2. Perbandingan throughput")]:
        path = CHARTS / image
        if path.exists():
            document.add_picture(str(path), width=Inches(6.1))
            add_caption(document, caption)
    add_heading(document, "4.5 Analisis", 2)
    document.add_paragraph("Distributed dua worker memberikan waktu tercepat pada pengujian 300 gambar. Percepatan belum linear karena seluruh worker berjalan pada satu node dengan dua CPU core. Overhead Redis, serialisasi task, scheduling Celery, dan I/O gambar ikut memengaruhi hasil. Temuan ini menunjukkan pentingnya menguji ukuran dataset yang cukup besar dan membandingkan jumlah worker dengan resource CPU yang tersedia.")
    add_heading(document, "4.6 Dokumentasi Screenshot", 2)
    document.add_paragraph("Screenshot terminal dapat ditempatkan pada bagian ini: spesifikasi server, Redis PONG, dua worker online, log task worker, master CLI, monitoring top, summary.json, comparison.csv, serta gambar input dan output.")

    add_heading(document, "BAB V PENUTUP", 1)
    add_heading(document, "5.1 Kesimpulan", 2)
    document.add_paragraph("Project berhasil menerapkan pemrosesan citra digital paralel dan terdistribusi menggunakan arsitektur master-worker, Redis, Celery, dan OpenCV pada Ubuntu Server. Sebanyak 300 gambar berhasil diproses tanpa task gagal. Konfigurasi distributed dua worker memberikan hasil terbaik dengan speedup 1,21 dan efficiency 60,74% pada eksperimen yang tercatat.")
    add_heading(document, "5.2 Saran", 2)
    for text in ["Melakukan pengujian pada beberapa server fisik.", "Menggunakan shared storage atau object storage untuk multi-node.", "Menambahkan dashboard Flower untuk monitoring.", "Menggunakan dataset citra nyata yang lebih beragam.", "Menguji GPU atau algoritma pemrosesan yang lebih berat."]:
        add_bullet(document, text)

    add_heading(document, "DAFTAR PUSTAKA", 1)
    for text in ["OpenCV Documentation. Image Processing Module.", "Celery Documentation. Distributed Task Queue.", "Redis Documentation. In-memory Data Store and Message Broker.", "AlmaLinux Documentation. Installation and System Administration."]:
        document.add_paragraph(text)

    document.add_heading("LAMPIRAN A — Struktur Project", level=1)
    document.add_paragraph("Repository: https://github.com/zenttzy/sistem-terdistribusi-image-processing")
    document.add_paragraph("Dokumentasi teks: docs/hasil-pengujian.txt")
    document.add_paragraph("Kode utama: app/cli.py, app/tasks.py, app/image_processor.py, app/reporting.py")
    document.add_paragraph("Grafik: reports/aggregate/execution-time.png dan reports/aggregate/throughput.png")

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
